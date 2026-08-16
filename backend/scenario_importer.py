"""剧本导入工具——读取常见剧本格式、切分文本、生成剧本总结与新剧本。

支持格式:
- .txt / .md / .markdown: 纯文本
- .pdf: pypdf 解析
- .docx: python-docx 解析
- .doc: 尽力解析（优先 antiword，其次原始文本提取）

切分策略:
- naive: 按段落/字数硬切分，速度快、无额外依赖
- semantic: 基于字符 n-gram 的局部语义相似度切分，适合长剧本保留语义连贯
"""

from __future__ import annotations

import io
import json
import math
import os
import re
import subprocess
from typing import Any

from openai import AsyncOpenAI

from backend.config import ensure_valid_api_key, settings
from backend.engine.game_systems import detect_game_system


# ═══════════════════════════════════════════════════════════════
# 文档读取
# ═══════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".doc"}


def extract_text(filename: str, data: bytes) -> str:
    """从上传文件内容中提取纯文本。"""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的剧本格式: {ext or '(无扩展名)'}，支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    if ext in (".txt", ".md", ".markdown"):
        return data.decode("utf-8", errors="replace")
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".doc":
        return _extract_doc(data)
    raise ValueError("无法读取该文件")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfStreamError
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n\n".join(pages)
        if not text.strip():
            raise ValueError("PDF 中没有可提取的文本（可能是扫描件，请改用文字版 PDF）")
        return text
    except (PdfStreamError, Exception) as e:
        raise ValueError(f"PDF 解析失败：文件可能损坏或不是有效 PDF（{e}）") from e


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError("DOCX 中没有可提取的文本")
    return text


def _extract_doc(data: bytes) -> str:
    """老式 .doc 是 OLE 复合格式，跨平台提取较麻烦。

    优先调用系统 antiword；不可用时退回提取可打印 ASCII/UTF-8 片段。
    """
    try:
        proc = subprocess.run(
            ["antiword", "-"],
            input=data,
            capture_output=True,
            timeout=20,
        )
        if proc.returncode == 0:
            text = proc.stdout.decode("utf-8", errors="replace").strip()
            if text:
                return text
    except Exception:
        pass

    # 退回：尝试 UTF-8 解码，失败则只保留可打印字符
    try:
        text = data.decode("utf-8", errors="strict")
    except UnicodeDecodeError:
        text = "".join(chr(b) for b in data if b in (9, 10, 13) or 32 <= b < 127)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise ValueError("无法从 .doc 中提取文本，请转换为 .docx / .txt / .pdf 后重试")
    return text


# ═══════════════════════════════════════════════════════════════
# 文本切分
# ═══════════════════════════════════════════════════════════════

def normalize_text(text: str) -> str:
    """清理文档中的常见噪声。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_text_naive(text: str, chunk_size: int = 900, overlap: int = 100) -> list[str]:
    """按段落和字数切分：快速、确定性强，适合大多数剧本。"""
    text = normalize_text(text)
    if not text:
        return []

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 1 <= chunk_size:
            current = f"{current}\n\n{para}".strip()
            continue
        # 当前块放不下新段：先按字符硬切当前块
        if current:
            chunks.append(current)
            current = current[-overlap:] if overlap else ""
        # 段落本身超过 chunk_size 时继续硬切
        while len(para) > chunk_size:
            chunks.append(para[:chunk_size])
            para = para[chunk_size - overlap:] if overlap else para[chunk_size:]
        current = para if not current else f"{current}\n\n{para}".strip()
    if current:
        chunks.append(current)
    return [c for c in chunks if c.strip()]


# ── 语义切分（纯 Python 实现，无需外部 embedding 服务） ──

def _char_ngrams(text: str, n: int = 2) -> list[str]:
    """提取字符 n-gram；中文按字，英文按词边界弱化。"""
    cleaned = re.sub(r"\s+", "", text.lower())
    if len(cleaned) <= n:
        return [cleaned] if cleaned else []
    return [cleaned[i:i + n] for i in range(len(cleaned) - n + 1)]


def _vec(text: str) -> dict[str, int]:
    vec: dict[str, int] = {}
    for gram in _char_ngrams(text):
        vec[gram] = vec.get(gram, 0) + 1
    return vec


def _cosine(a: dict[str, int], b: dict[str, int]) -> float:
    if not a or not b:
        return 0.0
    dot = 0
    for k, v in a.items():
        if k in b:
            dot += v * b[k]
    norm_a = math.sqrt(sum(v * v for v in a.values()))
    norm_b = math.sqrt(sum(v * v for v in b.values()))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def _merge_vec(base: dict[str, int], other: dict[str, int]) -> dict[str, int]:
    merged = dict(base)
    for k, v in other.items():
        merged[k] = merged.get(k, 0) + v
    return merged


def _join_chunk(sentences: list[str]) -> str:
    """按语言习惯拼接句子：中文无空格，英文/混合文本用空格分隔。"""
    if any(re.search(r"[\u4e00-\u9fff]", s) for s in sentences):
        return "".join(sentences)
    return " ".join(sentences)


def split_text_semantic(
    text: str,
    max_chunk_size: int = 1200,
    min_chunk_size: int = 400,
    similarity_threshold: float = 0.30,
) -> list[str]:
    """基于相邻句子的字符 n-gram 相似度进行语义切分。

    句子在主题发生明显变化（相似度低于阈值）时断开，且保证块长在
    [min_chunk_size, max_chunk_size] 附近；不会产生无意义的小碎片。
    """
    text = normalize_text(text)
    if not text:
        return []

    # 先按段落粗分，段内再按句子切，保留标点
    raw_sentences: list[str] = []
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        sentences = re.split(r"(?<=[。！？；;])|(?<=[.!?])\s+|\n", para)
        for s in sentences:
            s = s.strip()
            if s:
                raw_sentences.append(s)

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    current_vec: dict[str, int] | None = None

    def flush():
        nonlocal current, current_len, current_vec
        if current:
            chunks.append(_join_chunk(current))
        current = []
        current_len = 0
        current_vec = None

    for sent in raw_sentences:
        vec = _vec(sent)
        if current and current_len >= min_chunk_size:
            sim = _cosine(current_vec or {}, vec)
            if sim < similarity_threshold or current_len + len(sent) > max_chunk_size:
                flush()
        current.append(sent)
        current_len += len(sent)
        current_vec = vec if current_vec is None else _merge_vec(current_vec, vec)

    flush()
    # 极长块保护：语义切分结果中若仍有超过 max_chunk_size 的块，用硬切补齐
    final: list[str] = []
    for chunk in chunks:
        while len(chunk) > max_chunk_size * 1.5:
            final.append(chunk[:max_chunk_size])
            chunk = chunk[max_chunk_size - 100:]
        if chunk:
            final.append(chunk)
    return [c.strip() for c in final if c.strip()]


def split_text(text: str, mode: str = "naive", chunk_size: int = 900) -> list[str]:
    """对外切分入口。mode: naive | semantic"""
    if mode == "semantic":
        return split_text_semantic(text, max_chunk_size=max(600, chunk_size))
    return split_text_naive(text, chunk_size=chunk_size)


# ═══════════════════════════════════════════════════════════════
# LLM 剧本总结
# ═══════════════════════════════════════════════════════════════

SUMMARY_PROMPT = """你是一位TRPG模组编辑。请为以下冒险剧本写一份**约400字（允许350-450字）的剧本总结**。

总结要求：
1. 概括世界观、核心冲突、主线三幕、关键NPC、主要地点与独特规则
2. 语言精炼、信息密度高，让新玩家读完后能快速理解这是一个怎样的冒险
3. 使用中文，不要使用Markdown标题、列表符号，直接输出一段连贯文字
4. 不要输出"本剧本""总结如下"等多余引导语

## 冒险大纲
{outline}

## 原始剧本片段（可能截断）
{source}
"""


def _fallback_summary(outline: str, max_chars: int = 450) -> str:
    """从大纲中提取结构化摘要：每个章节取标题+首句；无标题时取每段首句，避免直接截断。"""
    if re.search(r"^#+\s", outline, re.M):
        lines = outline.split("\n")
        parts: list[str] = []
        current_heading = ""
        seen_heading_content = False
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                current_heading = stripped.lstrip("#").strip()
                if current_heading and (not parts or parts[-1] != current_heading):
                    parts.append(current_heading)
                seen_heading_content = False
                continue
            if current_heading and not seen_heading_content:
                sentence = re.split(r"(?<=[。！？!?；;])", stripped)[0][:100].strip()
                if sentence:
                    parts.append(sentence)
                    seen_heading_content = True
            if len("；".join(parts)) >= max_chars:
                break
        fallback = "；".join(parts)
    else:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", outline) if p.strip()]
        parts = []
        for para in paragraphs:
            sentence = re.split(r"(?<=[。！？!?；;])", para)[0][:100].strip()
            if sentence:
                parts.append(sentence)
            if len("；".join(parts)) >= max_chars:
                break
        fallback = "；".join(parts)
    fallback = fallback or outline.strip().replace("#", "").replace("*", "")
    fallback = re.sub(r"\s+", " ", fallback).strip()
    return fallback[:max_chars] or "（暂无剧本总结）"


async def generate_summary(
    client: AsyncOpenAI,
    model: str,
    outline: str,
    source_text: str,
    max_chars: int = 450,
) -> str:
    """调用 LLM 生成剧本总结；失败时使用结构化降级摘要，而不是简单截断。"""
    import asyncio
    last_err = None
    for attempt in range(1, 3):
        current_max_tokens = 2000 if attempt == 1 else 4000
        try:
            resp = await asyncio.wait_for(
                client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "你是一位严谨的TRPG模组编辑，擅长写出高信息密度的中文剧本总结。"},
                        {"role": "user", "content": SUMMARY_PROMPT.format(
                            outline=outline[:8000],
                            source=source_text[:4000],
                        )},
                    ],
                    max_tokens=current_max_tokens,
                    temperature=0.4,
                ),
                timeout=60,
            )
            summary = (resp.choices[0].message.content or "").strip()
            if len(summary) > max_chars * 1.4:
                summary = summary[:max_chars]
            if summary:
                return summary
            reasoning = getattr(resp.choices[0].message, "reasoning_content", None)
            print(f"[ScenarioImporter] 总结生成第{attempt}次空响应 (reasoning_len={len(reasoning or '')}, max_tokens={current_max_tokens})")
            raise RuntimeError("空响应")
        except Exception as e:
            last_err = str(e)
            print(f"[ScenarioImporter] 总结生成第{attempt}次失败: {e}")
        await asyncio.sleep(1)
    print(f"[ScenarioImporter] 总结生成最终失败: {last_err}，使用结构化降级摘要")
    return _fallback_summary(outline, max_chars)


# ═══════════════════════════════════════════════════════════════
# 从导入文本生成完整新剧本
# ═══════════════════════════════════════════════════════════════

async def generate_scenario_from_text(
    source_text: str,
    chunks: list[str],
    title: str = "",
    description: str = "",
    tone: str = "史诗奇幻",
    system: str | None = None,
    custom_rules: str = "",
    character_name: str = "冒险者",
    race: str = "人类",
    char_class: str = "战士",
    character_level: int = 1,
    api_key: str | None = None,
    model_name: str | None = None,
    base_url: str | None = None,
    splitter: str = "naive",
    target_score: int = 80,
    max_revisions: int = 2,
    custom_classes: list[str] | None = None,
    custom_skills: list[str] | None = None,
    extra_attributes: dict | None = None,
    thinking_strength: str = "medium",
) -> dict[str, Any]:
    """读取文本→切分→多Agent生成新剧本→生成总结→保存到 scenarios/。"""
    from backend.engine.world_builder import build_world
    from backend.scenario_store import create_scenario

    api_key = api_key or settings.LLM_API_KEY
    base_url = base_url or settings.LLM_BASE_URL
    model = model_name or settings.LLM_MODEL_NAME
    if not model:
        raise ValueError("请提供模型名称")
    api_key = ensure_valid_api_key(api_key)
    client = AsyncOpenAI(api_key=api_key, base_url=base_url)

    if not system or system == "auto":
        system = detect_game_system(source_text, title)

    player_input = (
        f"冒险基调: {tone}\n"
        f"规则系统: {system}\n"
        f"角色: {character_name}, {race} {char_class}, Lv.{character_level}\n"
        f"描述: {description or '根据导入的剧本生成一个完整冒险'}"
    )

    # 将切分后的块交给世界生成器；块之间用分隔符保留语义边界
    # 超长剧本先截断到安全长度，避免超出 LLM 上下文窗口
    reference_script = "\n\n===== 剧本片段 =====\n\n".join(
        f"[片段 {i + 1}/{len(chunks)}]\n{chunk}" for i, chunk in enumerate(chunks)
    ) if chunks else source_text
    if len(reference_script) > 30000:
        reference_script = reference_script[:30000] + "\n\n...[内容过长，已截断用于世界生成]..."

    outline_text, score, history, world_state = await build_world(
        player_input=player_input,
        reference_script=reference_script,
        api_key=api_key,
        model_name=model_name,
        base_url=base_url,
        game_system=system,
        custom_rules=custom_rules,
        custom_classes=custom_classes,
        custom_skills=custom_skills,
        extra_attributes=extra_attributes,
        target_score=target_score,
        max_revisions=max_revisions,
        thinking_strength=thinking_strength,
    )

    summary = await generate_summary(client, model, outline_text, source_text)

    ws_json = {
        "world_outline": outline_text,
        "npcs": [
            {
                "name": n.name, "race": n.race, "role": n.role,
                "location": n.location, "attitude": n.attitude,
                "alive": n.alive, "personality": n.personality,
                "motivation": n.motivation, "secret": n.secret,
                "relation_to_plot": n.relation_to_plot,
                "visibility": n.visibility.to_dict() if hasattr(n.visibility, "to_dict") else {},
            }
            for n in world_state.npcs
        ],
        "plot_flags": [
            {"key": f.key, "status": f.status, "description": f.description}
            for f in world_state.plot_flags
        ],
        "locations": [
            {"name": l.name, "description": l.description, "secrets": l.secrets}
            for l in world_state.locations
        ],
        "world_rules": world_state.world_rules,
    }
    world_state_json = json.dumps(ws_json, ensure_ascii=False)

    saved = create_scenario(
        world_outline=outline_text,
        world_state_json=world_state_json,
        reference_script=source_text,
        source_chunks=chunks,
        custom_rules=custom_rules,
        custom_classes=custom_classes or [],
        custom_skills=custom_skills or [],
        extra_attributes=extra_attributes or {},
        notes=f"导入方式: {splitter} 切分 · 共 {len(chunks)} 个片段",
        title=title or (outline_text.split("\n")[0].replace("#", "").strip()[:60] or "导入冒险"),
        description=description or source_text[:200],
        summary=summary,
        system=system,
        tone=tone,
        character_name=character_name,
        race=race,
        char_class=char_class,
        level=character_level,
        score=score,
    )

    # 将剧本细节写入本地知识库，供后续 RAG 检索
    try:
        from backend.knowledge_base import get_knowledge_base
        kb = get_knowledge_base()
        scenario_source = f"scenario:{saved.id}"
        for d in kb.list_documents():
            if d.get("source") == scenario_source:
                kb.remove_document(d["id"])
        kb.add_document(
            title=f"剧本：{saved.meta.title}",
            content=source_text,
            source=scenario_source,
            system=system,
            tags=["剧本", system, splitter],
        )
    except Exception as e:
        print(f"[ScenarioImporter] 知识库写入失败（不影响剧本生成）: {e}")

    npc_summary = [{"name": n.name, "role": n.role, "attitude": n.attitude}
                   for n in world_state.npcs]
    flag_summary = [{"key": f.key, "status": f.status} for f in world_state.plot_flags]

    return {
        "scenario_id": saved.id,
        "title": saved.meta.title,
        "summary": summary,
        "system": saved.meta.system,
        "content": outline_text,
        "score": score,
        "scores_detail": {},
        "revision_history": history,
        "npcs": npc_summary,
        "plot_flags": flag_summary,
        "world_rules": world_state.world_rules,
        "world_state_json": world_state_json,
        "source_chunks": chunks,
        "chunk_count": len(chunks),
        "splitter": splitter,
    }
